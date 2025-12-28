import streamlit as st
import json
import re
import stripe
import smtplib
import time
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.application import MIMEApplication
from datetime import datetime
from io import BytesIO
import zipfile
import google.generativeai as genai
from pypdf import PdfReader
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF

# --- GESTIONE IMPORT CONDIZIONALE SUPABASE ---
try:
    from supabase import create_client, Client
    SUPABASE_AVAILABLE = True
except ImportError:
    SUPABASE_AVAILABLE = False

# ==============================================================================
# 1. CONFIGURAZIONE SISTEMA & UI
# ==============================================================================

APP_NAME = "LexVantage"
APP_VER = "Rev 50.3 (Monolith: Full SaaS Architecture)"
st.set_page_config(page_title=APP_NAME, layout="wide", page_icon="⚖️")

# CSS PREMIUM AVANZATO
st.markdown("""
<style>
    /* Stile Messaggi Chat */
    div[data-testid="stChatMessage"] { 
        background-color: #f8f9fa; 
        border-radius: 8px; 
        padding: 15px; 
        border-left: 5px solid #004e92; 
        margin-bottom: 10px;
        font-family: 'Segoe UI', sans-serif;
        box-shadow: 0 2px 4px rgba(0,0,0,0.05);
    }
    
    /* Box Paywall Giallo */
    .premium-box { 
        padding: 30px; 
        border: 2px solid #ffd700; 
        background-color: #fffdf0; 
        border-radius: 12px; 
        text-align: center; 
        margin-top: 20px; 
        box-shadow: 0 4px 12px rgba(0,0,0,0.1); 
    }
    
    /* Box Login */
    .auth-box { 
        max-width: 450px; 
        margin: 50px auto; 
        padding: 30px; 
        border: 1px solid #e0e0e0; 
        border-radius: 15px; 
        background: white; 
        box-shadow: 0 8px 20px rgba(0,0,0,0.1);
    }
    
    /* Tipografia */
    h1, h2, h3 { color: #003366; font-family: 'Helvetica Neue', sans-serif; }
    
    /* Bottoni */
    .stButton>button { 
        width: 100%; 
        font-weight: 600; 
        border-radius: 8px; 
        height: 3rem;
        transition: transform 0.1s;
    }
    .stButton>button:active { transform: scale(0.98); }
</style>
""", unsafe_allow_html=True)

# ==============================================================================
# 2. INIZIALIZZAZIONE SERVIZI (DB, AI, EMAIL, PAYMENT)
# ==============================================================================

@st.cache_resource
def init_supabase():
    """Inizializza la connessione al database Supabase"""
    if not SUPABASE_AVAILABLE: return None
    try:
        url = st.secrets["supabase"]["url"]
        key = st.secrets["supabase"]["key"]
        return create_client(url, key)
    except Exception as e:
        return None

supabase = init_supabase()

# Inizializzazione Session State Completo
if "auth_status" not in st.session_state: st.session_state.auth_status = "logged_out"
if "user_role" not in st.session_state: st.session_state.user_role = "user"
if "user_email" not in st.session_state: st.session_state.user_email = ""
if "user_id" not in st.session_state: st.session_state.user_id = None
if "messages" not in st.session_state: st.session_state.messages = []
if "contesto_chat" not in st.session_state: st.session_state.contesto_chat = ""
if "dati_calc" not in st.session_state: st.session_state.dati_calc = "Nessun calcolo effettuato."
if "generated_docs" not in st.session_state: st.session_state.generated_docs = {}
if "workflow_step" not in st.session_state: st.session_state.workflow_step = "CHAT"
if "token_cost" not in st.session_state: st.session_state.token_cost = 150.00 

# Configurazione API Esterne
HAS_KEY = False
active_model = None
PAYMENT_ENABLED = False
EMAIL_ENABLED = False

# AUTO-DISCOVERY MODELLO (Logica Rev 48 Ripristinata)
try:
    lista_modelli = [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
    
    # Ordine di preferenza: CERCA PRIMA IL PRO (Più intelligente), POI IL FLASH
    if "models/gemini-1.5-pro" in lista_modelli: 
        active_model = "models/gemini-1.5-pro"
    elif "models/gemini-1.5-pro-latest" in lista_modelli: 
        active_model = "models/gemini-1.5-pro-latest"
    elif "models/gemini-1.5-flash" in lista_modelli: 
        active_model = "models/gemini-1.5-flash"
    else:
        active_model = lista_modelli[0]
        
    print(f"Modello AI selezionato: {active_model}")

except Exception as e:
    active_model = "models/gemini-1.5-flash" # Fallback estremo

# ==============================================================================
# 3. CORE LOGIC: PRIVACY & FORMATTAZIONE (DALLA REV 48)
# ==============================================================================

class DataSanitizer:
    """Gestisce l'anonimizzazione dei dati sensibili (GDPR Compliance)"""
    def __init__(self):
        self.mapping = {}
        self.reverse = {}
        self.cnt = 1
        
    def add(self, real, label):
        if real and real not in self.mapping:
            fake = f"[{label}_{self.cnt}]"
            self.mapping[real] = fake
            self.reverse[fake] = real
            self.cnt += 1
            
    def sanitize(self, txt):
        if not txt: return ""
        for r, f in self.mapping.items():
            txt = txt.replace(r, f).replace(r.upper(), f)
        return txt
        
    def restore(self, txt):
        if not txt: return ""
        for f, r in self.reverse.items():
            txt = txt.replace(f, r)
        return txt

if "sanitizer" not in st.session_state: st.session_state.sanitizer = DataSanitizer()

def universal_json_flattener(data, level=0):
    """Converte JSON nidificati complessi in testo Markdown leggibile"""
    text = ""
    indent = "  " * level
    if isinstance(data, dict):
        if "titolo" in data and "contenuto" in data: 
            return f"### {data['titolo']}\n\n{universal_json_flattener(data['contenuto'], level)}"
        for k, v in data.items():
            if isinstance(v, (dict, list)): 
                text += f"\n{indent}**{k.title()}**:\n{universal_json_flattener(v, level+1)}"
            else: 
                text += f"\n{indent}- **{k.title()}**: {v}"
    elif isinstance(data, list):
        for item in data: 
            text += f"\n{indent}* {universal_json_flattener(item, level+1)}"
    else: 
        return str(data).replace("|", " - ") # Previene rottura tabelle Markdown
    return text.strip()

def parse_markdown_pro(doc, text):
    """Motore di rendering DOCX: Converte Markdown in Word con stili reali"""
    lines = str(text).split('\n')
    in_table = False
    table_data = []
    
    for line in lines:
        stripped = line.strip()
        
        # Rilevamento Tabelle
        if "|" in stripped and stripped.startswith("|"):
            if not in_table: in_table=True; table_data=[]
            cells = [c.strip() for c in stripped.split('|')[1:-1]]
            # Ignora righe separatore markdown (es: |---|---|)
            if not all(set(c).issubset({'-', ':', ' '}) for c in cells if c): 
                table_data.append(cells)
            continue
            
        # Rendering Tabella
        if in_table:
            if table_data:
                rows = len(table_data)
                cols = max(len(r) for r in table_data) if rows>0 else 0
                if rows>0 and cols>0:
                    tbl = doc.add_table(rows, cols)
                    tbl.style = 'Table Grid'
                    for i,r in enumerate(table_data):
                        for j,c in enumerate(r): 
                            if j<cols: tbl.cell(i,j).text=c
            in_table=False
            table_data=[]
            
        if not stripped: continue
        
        # Rendering Intestazioni ed Elenchi
        if stripped.startswith('#'): 
            level = min(stripped.count('#'), 3)
            doc.add_heading(stripped.lstrip('#').strip(), level=level)
        elif stripped.startswith('- ') or stripped.startswith('* '): 
            doc.add_paragraph(stripped[2:], style='List Bullet')
        else: 
            p = doc.add_paragraph(stripped)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# ==============================================================================
# 4. SISTEMA EMAIL (NOTIFICATION & APPROVAL)
# ==============================================================================

def send_admin_alert(new_user_email):
    """Invia email all'Admin quando un nuovo utente si registra"""
    if not EMAIL_ENABLED: return
    try:
        msg = MIMEMultipart()
        msg['Subject'] = f"🔔 Nuovo Iscritto LexVantage: {new_user_email}"
        msg['From'] = st.secrets["smtp"]["email"]
        msg['To'] = st.secrets["smtp"]["email"] # Invia a se stesso
        
        body = f"""
        Un nuovo utente ha richiesto l'accesso.
        Email: {new_user_email}
        
        Accedi al Pannello Admin per approvare o rifiutare.
        """
        msg.attach(MIMEText(body, 'plain'))
        
        s = smtplib.SMTP(st.secrets["smtp"]["server"], st.secrets["smtp"]["port"])
        s.starttls()
        s.login(st.secrets["smtp"]["email"], st.secrets["smtp"]["password"])
        s.sendmail(st.secrets["smtp"]["email"], st.secrets["smtp"]["email"], msg.as_string())
        s.quit()
    except Exception as e:
        print(f"Errore invio mail admin: {e}")

def send_approval_email(user_email):
    """Invia email all'Utente quando viene approvato"""
    if not EMAIL_ENABLED: return False
    try:
        msg = MIMEMultipart()
        msg['Subject'] = "✅ Benvenuto in LexVantage - Account Attivo"
        msg['From'] = st.secrets["smtp"]["email"]
        msg['To'] = user_email
        
        body = """
        Gentile Utente,
        
        Siamo lieti di informarti che il tuo account LexVantage è stato approvato dall'Amministratore.
        
        Puoi ora accedere alla piattaforma utilizzando la tua email e la password scelta in fase di registrazione.
        
        Buon lavoro,
        Il Team LexVantage
        """
        msg.attach(MIMEText(body, 'plain'))
        
        s = smtplib.SMTP(st.secrets["smtp"]["server"], st.secrets["smtp"]["port"])
        s.starttls()
        s.login(st.secrets["smtp"]["email"], st.secrets["smtp"]["password"])
        s.sendmail(st.secrets["smtp"]["email"], user_email, msg.as_string())
        s.quit()
        return True
    except Exception as e:
        print(f"Errore invio mail utente: {e}")
        return False

# ==============================================================================
# 5. MODULI DI AUTENTICAZIONE E AMMINISTRAZIONE
# ==============================================================================

def login_form():
    """Gestisce Login e Registrazione con DB Supabase Custom"""
    st.markdown("<div class='auth-box'>", unsafe_allow_html=True)
    st.markdown("<h2 style='text-align: center;'>🔐 Accedi a LexVantage</h2>", unsafe_allow_html=True)
    
    tab_login, tab_register = st.tabs(["Login", "Registrati"])
    
    with tab_login:
        email = st.text_input("Email", key="log_email")
        pwd = st.text_input("Password", type="password", key="log_pwd")
        
        if st.button("Accedi", type="primary"):
            if supabase:
                # Query sulla tabella profili_utenti
                res = supabase.table("profili_utenti").select("*").eq("email", email).eq("password", pwd).execute()
                if res.data:
                    user = res.data[0]
                    if user['stato_account'] == 'attivo':
                        st.session_state.auth_status = "logged_in"
                        st.session_state.user_email = user['email']
                        st.session_state.user_role = user['ruolo']
                        st.session_state.user_id = user['id']
                        st.rerun()
                    elif user['stato_account'] == 'in_attesa':
                        st.warning("⏳ Il tuo account è in attesa di approvazione.")
                    else:
                        st.error("🚫 Account sospeso o disattivato.")
                else:
                    st.error("Email o password errati.")
            else:
                # Fallback Offline (Solo se DB non connesso)
                if email=="admin" and pwd=="admin":
                    st.session_state.auth_status="logged_in"; st.session_state.user_role="admin"; st.rerun()
                elif email=="user" and pwd=="user":
                    st.session_state.auth_status="logged_in"; st.session_state.user_role="user"; st.rerun()

    with tab_register:
        st.write("Crea un nuovo account per il tuo Studio.")
        reg_email = st.text_input("Email", key="reg_email")
        reg_pwd = st.text_input("Password", type="password", key="reg_pwd")
        reg_studio = st.text_input("Nome Studio Legale", key="reg_studio")
        
        if st.button("Richiedi Accesso"):
            if reg_email and reg_pwd and reg_studio and supabase:
                try:
                    # Verifica esistenza
                    check = supabase.table("profili_utenti").select("email").eq("email", reg_email).execute()
                    if check.data:
                        st.error("Email già registrata.")
                    else:
                        supabase.table("profili_utenti").insert({
                            "email": reg_email, 
                            "password": reg_pwd, 
                            "nome_studio": reg_studio, 
                            "ruolo": "user",
                            "stato_account": "in_attesa"
                        }).execute()
                        
                        # Invia avviso all'admin
                        send_admin_alert(reg_email)
                        
                        st.success("✅ Richiesta inviata con successo! Riceverai una mail appena l'account sarà attivo.")
                except Exception as e:
                    st.error(f"Errore durante la registrazione: {str(e)}")
            else:
                st.warning("Compila tutti i campi.")
                
    st.markdown("</div>", unsafe_allow_html=True)

def admin_panel():
    """Dashboard completa per l'Amministratore"""
    st.markdown("## 🛠️ Admin Dashboard")
    st.info(f"Loggato come Superuser: {st.session_state.user_email}")
    
    tab_users, tab_prices, tab_audit = st.tabs(["👥 Gestione Utenti", "💰 Listino Prezzi", "📂 Audit Fascicoli"])
    
    # --- 1. GESTIONE UTENTI ---
    with tab_users:
        st.subheader("Richieste di Accesso in Attesa")
        if supabase:
            pending = supabase.table("profili_utenti").select("*").eq("stato_account", "in_attesa").execute().data
            
            if not pending:
                st.success("Nessuna richiesta in sospeso.")
            
            for u in pending:
                with st.container():
                    col_info, col_act = st.columns([3, 1])
                    with col_info:
                        st.markdown(f"**{u['nome_studio']}** ({u['email']})")
                        st.caption(f"Richiesto il: {u['created_at']}")
                    with col_act:
                        if st.button("✅ APPROVA", key=f"app_{u['id']}"):
                            # Attivazione
                            supabase.table("profili_utenti").update({"stato_account": "attivo"}).eq("id", u['id']).execute()
                            # Notifica Email
                            with st.spinner("Invio email di conferma..."):
                                res_mail = send_approval_email(u['email'])
                            
                            if res_mail: st.toast("Utente attivato e avvisato via mail!", icon="📧")
                            else: st.toast("Utente attivato, ma errore invio mail.", icon="⚠️")
                            
                            time.sleep(2)
                            st.rerun()
                            
                        if st.button("❌ RIFIUTA", key=f"den_{u['id']}"):
                            supabase.table("profili_utenti").delete().eq("id", u['id']).execute()
                            st.warning("Utente rimosso.")
                            time.sleep(1)
                            st.rerun()
                    st.divider()
        else:
            st.error("Database non connesso.")

    # --- 2. GESTIONE PREZZI ---
    with tab_prices:
        st.subheader("Configurazione Dinamica Prezzi")
        if supabase:
            try:
                prezzi = supabase.table("listino_prezzi").select("*").execute().data
                for p in prezzi:
                    with st.form(key=f"form_price_{p['id']}"):
                        st.markdown(f"### {p['tipo_documento'].replace('_', ' ').upper()}")
                        c1, c2, c3 = st.columns(3)
                        
                        new_fixed = c1.number_input("Prezzo Fisso (€)", value=float(p['prezzo_fisso']))
                        new_token_in = c2.number_input("Costo Input / 1k Token (€)", value=float(p['prezzo_per_1k_input_token']), format="%.4f")
                        new_token_out = c3.number_input("Costo Output / 1k Token (€)", value=float(p['prezzo_per_1k_output_token']), format="%.4f")
                        
                        if st.form_submit_button("💾 Salva Modifiche"):
                            supabase.table("listino_prezzi").update({
                                "prezzo_fisso": new_fixed,
                                "prezzo_per_1k_input_token": new_token_in,
                                "prezzo_per_1k_output_token": new_token_out
                            }).eq("id", p['id']).execute()
                            st.success(f"Listino per {p['tipo_documento']} aggiornato!")
                            time.sleep(1)
                            st.rerun()
            except Exception as e:
                st.error(f"Errore caricamento prezzi: {e}")

    # --- 3. AUDIT FASCICOLI ---
    with tab_audit:
        st.subheader("Monitoraggio Globale Fascicoli")
        if supabase:
            # Recupera ultimi 50 fascicoli
            fascicoli = supabase.table("fascicoli").select("*").order("created_at", desc=True).limit(50).execute().data
            if fascicoli:
                st.dataframe(fascicoli)
            else:
                st.info("Nessun fascicolo presente nel sistema.")

# ==============================================================================
# 6. FUNZIONI AI & DOCUMENT PROCESSING
# ==============================================================================

def get_file_content(uploaded_files):
    """Estrae testo da PDF, DOCX e Immagini"""
    parts = []
    full_text = ""
    
    parts.append("FASCICOLO DOCUMENTALE CARICATO:\n")
    
    if not uploaded_files: 
        return [], ""
        
    for file in uploaded_files:
        try:
            filename_clean = file.name.replace("|", "_")
            text_extracted = ""
            
            # PDF Processing
            if file.type == "application/pdf":
                reader = PdfReader(file)
                text_extracted = "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
                
            # Word Processing
            elif "word" in file.type:
                doc = Document(file)
                text_extracted = "\n".join([para.text for para in doc.paragraphs])
            
            # Text Processing
            elif "text" in file.type:
                text_extracted = str(file.read(), "utf-8")
                
            if text_extracted:
                header = f"\n--- INIZIO DOCUMENTO: {filename_clean} ---\n"
                parts.append(header + text_extracted)
                full_text += header + text_extracted
                
        except Exception as e:
            st.error(f"Errore lettura file {file.name}: {str(e)}")
            
    return parts, full_text

def interroga_gemini(prompt, history, files_parts, is_commit_phase=False):
    """Chiamata principale all'LLM con contesto e sanitizzazione"""
    if not HAS_KEY: return "⚠️ ERRORE: Chiave AI non configurata."
    
    sanitizer = st.session_state.sanitizer
    safe_prompt = sanitizer.sanitize(prompt)
    safe_history = sanitizer.sanitize(history)
    
    # Costruzione System Prompt
    mood = "AGGRESSIVO, PRAGMATICO, ORIENTATO AL RISULTATO ECONOMICO"
    sys_prompt = f"""
    SEI LEXVANTAGE, un assistente legale strategico di alto livello.
    MOOD: {mood}.
    
    CONTESTO TECNICO DEL CALCOLATORE:
    {st.session_state.dati_calc}
    
    STORICO CONVERSAZIONE:
    {safe_history}
    """
    
    if is_commit_phase:
        sys_prompt += """
        \nTASK SPECIALE (CHIUSURA):
        L'utente ha chiesto di generare i documenti finali.
        NON GENERARE ANCORA I DOCUMENTI COMPLETI.
        Invece, fai una sintesi della strategia emersa e chiedi conferma esplicita.
        
        DEVI RISPONDERE CON QUALCOSA DI SIMILE A:
        'Ho analizzato il fascicolo. La strategia vincente è [SINTESI]. 
        Il costo per generare il pacchetto documenti completo è di € [PREZZO].
        Confermi di voler procedere all'acquisto?'
        """
    else:
        sys_prompt += """
        \nTASK STANDAD:
        Rispondi alla domanda dell'utente basandoti sui documenti.
        Se l'utente chiede 'creami il documento', rispondi che puoi farlo nel Tab 'Documenti' dopo aver definito la strategia.
        Sii sintetico e diretto.
        """
        
    payload = list(files_parts)
    payload.append(f"UTENTE: {safe_prompt}")
    
    try:
        model = genai.GenerativeModel(active_model, system_instruction=sys_prompt)
        response = model.generate_content(payload)
        return sanitizer.restore(response.text)
    except Exception as e:
        return f"Errore AI: {str(e)}"

def genera_docs_json_batch(tasks, context, file_parts):
    """Generazione massiva di documenti strutturati JSON"""
    results = {}
    
    model = genai.GenerativeModel(
        active_model, 
        generation_config={"response_mime_type": "application/json"}
    )
    
    for doc_name, task_instruction in tasks:
        full_payload = list(file_parts)
        
        specific_prompt = f"""
        CONTESTO GENERALE:
        {context}
        
        OBIETTIVO SPECIFICO DOCUMENTO '{doc_name}':
        {task_instruction}
        
        FORMATO OUTPUT RICHIESTO:
        JSON rigoroso con due chiavi: "titolo" e "contenuto".
        Il campo "contenuto" deve essere in formato MARKDOWN ricco (grassetti, elenchi puntati, tabelle |...|).
        """
        
        full_payload.append(specific_prompt)
        
        try:
            raw_response = model.generate_content(full_payload).text
            parsed_json = json.loads(raw_response)
            results[doc_name] = parsed_json
        except Exception as e:
            results[doc_name] = {
                "titolo": f"Errore {doc_name}",
                "contenuto": f"Impossibile generare il documento. Dettagli: {str(e)}"
            }
            
    return results

# ==============================================================================
# 7. MAIN APPLICATION FLOW
# ==============================================================================

if st.session_state.auth_status != "logged_in":
    login_form()
else:
    # --- SIDEBAR UTENTE ---
    with st.sidebar:
        st.title(APP_NAME)
        st.caption(APP_VER)
        st.write(f"👤 **{st.session_state.user_email}**")
        
        # Accesso Admin
        if st.session_state.user_role == "admin":
            st.divider()
            if st.checkbox("🔧 Pannello Admin"):
                admin_panel()
                st.stop() # Blocca il rendering del resto se siamo in admin mode
        
        st.divider()
        if st.button("Esci (Logout)"):
            st.session_state.auth_status = "logged_out"
            st.rerun()
            
        st.divider()
        st.write("### 🛡️ Privacy Shield")
        st.caption("Maschera i nomi reali prima di inviarli all'AI.")
        n1 = st.text_input("Nome Cliente (Reale)", "Rossi")
        n2 = st.text_input("Controparte (Reale)", "Bianchi")
        
        if st.button("Attiva Mascheramento"):
            st.session_state.sanitizer.add(n1, "CLIENTE_PROTETTO")
            st.session_state.sanitizer.add(n2, "CONTROPARTE_PROTETTA")
            st.success("Privacy Shield Attivo!")

    # --- MAIN TABS ---
    t1, t2, t3 = st.tabs(["🧮 1. Calcolatore Tecnico", "💬 2. Analisi & Strategia", "📦 3. Generazione & Acquisto"])

    # TAB 1: CALCOLATORE
    with t1:
        st.header("Calcolatore Valore & Vizi")
        col1, col2 = st.columns(2)
        
        with col1:
            val_ctu = st.number_input("Valore Stimato CTU (€)", value=0.0, step=1000.0)
            val_target = st.number_input("Valore Target Nostro (€)", value=0.0, step=1000.0)
        
        with col2:
            st.markdown("#### Note Tecniche / Vizi")
            note_tecniche = st.text_area(
                "Inserisci dettagli su abusi, vizi strutturali, stato occupativo...",
                height=150
            )
            
        if st.button("💾 Salva Dati nel Contesto AI", type="primary"):
            st.session_state.dati_calc = f"""
            DATI ECONOMICI:
            - Valore CTU: € {val_ctu}
            - Valore Target: € {val_target}
            - Delta: € {val_ctu - val_target}
            
            NOTE TECNICHE:
            {note_tecniche}
            """
            st.success("Dati tecnici salvati e pronti per l'analisi.")

    # TAB 2: CHAT (LEAD MAGNET)
    with t2:
        st.subheader("Analisi Strategica del Fascicolo")
        
        uploaded_files = st.file_uploader("Carica documenti (PDF, DOCX, TXT)", accept_multiple_files=True)
        files_parts, full_text = get_file_content(uploaded_files)
        
        # Rendering Chat History
        for msg in st.session_state.messages:
            with st.chat_message(msg["role"]):
                st.markdown(msg["content"])
        
        # Input Chat
        if prompt := st.chat_input("Chiedi a LexVantage (es: Analizza la strategia avversaria)"):
            # 1. Aggiungi User Message
            st.session_state.messages.append({"role": "user", "content": prompt})
            st.session_state.contesto_chat += f"\nUTENTE: {prompt}"
            with st.chat_message("user"):
                st.write(prompt)
            
            # 2. Rileva Intento di Chiusura/Acquisto
            keywords_closure = ["genera", "documenti", "chiudi", "ordine", "procediamo", "comprare", "acquisto"]
            is_closing = any(k in prompt.lower() for k in keywords_closure)
            
            # 3. Chiamata AI
            with st.spinner("Analisi giuridica in corso..."):
                reply = interroga_gemini(
                    prompt, 
                    st.session_state.contesto_chat, 
                    files_parts, 
                    is_commit_phase=is_closing
                )
                
                # Se l'AI propone il prezzo, passiamo allo step COMMIT
                if "€" in reply and "?" in reply and is_closing:
                    st.session_state.workflow_step = "COMMIT"
                
                st.session_state.messages.append({"role": "assistant", "content": reply})
                st.session_state.contesto_chat += f"\nAI: {reply}"
                
                with st.chat_message("assistant"):
                    st.markdown(reply)
        
        # Bottoni di Conferma (Workflow Commit)
        if st.session_state.workflow_step == "COMMIT":
            st.divider()
            st.info("👆 L'AI ha proposto una strategia e un preventivo. Come vuoi procedere?")
            col_yes, col_no = st.columns(2)
            
            if col_yes.button("✅ Sì, Procedi all'Acquisto"):
                st.session_state.workflow_step = "PAYMENT"
                st.rerun()
                
            if col_no.button("❌ No, ho altre domande"):
                st.session_state.workflow_step = "CHAT"
                st.rerun()

    # TAB 3: PAYWALL & GENERAZIONE
    with t3:
        st.header("Generazione Fascicolo Esecutivo")
        
        # STEP 1: LOCKED
        if st.session_state.workflow_step == "CHAT":
            st.warning("⚠️ Completa l'analisi preliminare nel Tab 2 prima di generare i documenti.")
            st.markdown("""
            <div style="filter: blur(5px); opacity: 0.5;">
                <h3>Sintesi Esecutiva</h3>
                <p>Il documento contiene la strategia vincente...</p>
                <hr>
                <h3>Nota Difensiva</h3>
                <p>Contestazione puntuale delle perizie avversarie...</p>
            </div>
            """, unsafe_allow_html=True)
            
        # STEP 2: PAYMENT
        elif st.session_state.workflow_step == "PAYMENT":
            
            # Recupera prezzo da DB se disponibile
            current_price = st.session_state.token_cost
            if supabase:
                try:
                    p_data = supabase.table("listino_prezzi").select("prezzo_fisso").eq("tipo_documento", "pacchetto_base").execute()
                    if p_data.data:
                        current_price = p_data.data[0]['prezzo_fisso']
                except: pass
            
            st.markdown(f"""
            <div class='premium-box'>
                <h2>🚀 Sblocca il Fascicolo Completo</h2>
                <p>Hai confermato la strategia. Ora genera i documenti esecutivi pronti per il deposito.</p>
                <h1 style='color: #28a745; font-size: 3em;'>€ {current_price}</h1>
                <p>Il pacchetto include: Sintesi, Strategia, Note Difensive, Matrice Rischi, Quesiti CTU.</p>
            </div>
            """, unsafe_allow_html=True)
            
            st.divider()
            col_pay, col_sim = st.columns(2)
            
            with col_pay:
                if PAYMENT_ENABLED:
                    st.write("Pagamento Sicuro:")
                    # Qui andrebbe la logica stripe.checkout.Session.create reale
                    # Per ora linkiamo a una pagina di test o bottone
                    if st.button("💳 PAGA CON CARTA (Stripe)"):
                        with st.spinner("Connessione al gateway bancario..."):
                            time.sleep(1.5)
                        st.session_state.workflow_step = "UNLOCKED"
                        st.balloons()
                        st.rerun()
                else:
                    st.info("Sistema Pagamenti Stripe non configurato.")
            
            with col_sim:
                st.write("Debug Mode:")
                if st.button("🛠️ Simula Pagamento Riuscito"):
                    st.session_state.workflow_step = "UNLOCKED"
                    st.balloons()
                    st.rerun()

        # STEP 3: GENERATION (UNLOCKED)
        elif st.session_state.workflow_step == "UNLOCKED":
            st.success("✅ PAGAMENTO RICEVUTO. GENERAZIONE ABILITATA.")
            
            docs_map = {
                "Sintesi": "Sintesi Esecutiva e Strategica (per il Cliente)",
                "Nota_Difensiva": "Nota Difensiva Aggressiva (per il Giudice)",
                "Strategia": "Piano d'Azione Scenario A/B (Interno Studio)",
                "Quesiti_CTU": "Quesiti Tecnici Demolitori per il CTU",
                "Matrice_Rischi": "Matrice dei Rischi Economici",
                "Bozza_Transazione": "Bozza Accordo Transattivo"
            }
            
            selected_docs = st.multiselect(
                "Seleziona i documenti da generare:",
                list(docs_map.keys()),
                default=list(docs_map.keys())
            )
            
            if st.button("🚀 AVVIA GENERAZIONE FASCICOLO", type="primary"):
                if not uploaded_files:
                    st.warning("Attenzione: Stai generando senza file caricati nel Tab 2.")
                
                # Prepara i task
                tasks = [(k, docs_map[k]) for k in selected_docs]
                
                # Barra di progresso
                progress_bar = st.progress(0)
                status_text = st.empty()
                
                generated_results = {}
                
                # Ciclo di generazione
                total = len(tasks)
                for i, (doc_key, doc_prompt) in enumerate(tasks):
                    status_text.text(f"Scrittura in corso: {doc_key}...")
                    
                    # Generazione singola
                    res = genera_docs_json_batch(
                        [(doc_key, doc_prompt)], 
                        st.session_state.contesto_chat, 
                        files_parts
                    )
                    generated_results.update(res)
                    
                    progress_bar.progress((i + 1) / total)
                
                status_text.text("Compilazione file DOCX e ZIP...")
                
                # Creazione ZIP
                zip_buffer = BytesIO()
                with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
                    for doc_name, doc_data in generated_results.items():
                        # Crea DOCX
                        docx_obj = Document()
                        # Titolo
                        docx_obj.add_heading(doc_data.get("titolo", doc_name), 0)
                        # Contenuto (Restore Privacy + Parse Markdown)
                        clean_content = st.session_state.sanitizer.restore(str(doc_data.get("contenuto", "")))
                        parse_markdown_pro(docx_obj, clean_content)
                        
                        # Salva in buffer
                        docx_bytes = BytesIO()
                        docx_obj.save(docx_bytes)
                        
                        # Aggiungi a ZIP
                        zf.writestr(f"{doc_name}.docx", docx_bytes.getvalue())
                
                st.session_state.generated_docs = zip_buffer
                st.success("Tutti i documenti sono stati generati!")
                
            # Download Area
            if st.session_state.generated_docs:
                st.divider()
                st.download_button(
                    label="📦 SCARICA FASCICOLO COMPLETO (ZIP)",
                    data=st.session_state.generated_docs.getvalue(),
                    file_name="Fascicolo_LexVantage_Executive.zip",
                    mime="application/zip",
                    type="primary"
                )


