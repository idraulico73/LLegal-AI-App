# modules/doc_renderer.py
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pypdf import PdfReader
from io import BytesIO
import zipfile
import re

def extract_text_from_files(uploaded_files):
    """
    Estrae testo da PDF, DOCX e TXT caricati.
    Restituisce: (lista_parti, testo_completo)
    """
    parts = []
    full_text = ""
    parts.append("FASCICOLO DOCUMENTALE:\n")
    
    if not uploaded_files: 
        return [], ""
    
    for file in uploaded_files:
        try:
            txt = ""
            if file.type == "application/pdf":
                reader = PdfReader(file)
                # Estrae testo da tutte le pagine se presente
                txt = "\n".join([p.extract_text() for p in reader.pages if p.extract_text()])
            elif "word" in file.type or "docx" in file.name:
                doc = Document(file)
                txt = "\n".join([p.text for p in doc.paragraphs])
            else:
                # Fallback per file testo
                txt = str(file.read(), "utf-8")
                
            # Aggiungiamo separatori chiari per l'AI
            header = f"\n--- DOCUMENTO CARICATO: {file.name} ---\n"
            parts.append(f"{header}{txt}")
            full_text += f"{header}{txt}"
        except Exception as e: 
            parts.append(f"Errore lettura file {file.name}: {str(e)}")
            
    return parts, full_text

def parse_markdown_pro(doc, text):
    """
    Converte Markdown (tabelle, grassetti, titoli) in elementi nativi Word.
    """
    lines = str(text).split('\n')
    in_table = False
    table_data = []
    
    for line in lines:
        stripped = line.strip()
        
        # 1. Gestione Tabelle Markdown (| col | col |)
        if "|" in stripped and stripped.startswith("|"):
            if not in_table: in_table=True; table_data=[]
            cells = [c.strip() for c in stripped.split('|')[1:-1]]
            # Salta le righe di separazione (es. |---|---|)
            if not all(set(c).issubset({'-', ':', ' '}) for c in cells if c): 
                table_data.append(cells)
            continue
            
        if in_table:
            # Fine tabella rilevata, renderizziamo
            if table_data:
                rows = len(table_data)
                cols = max(len(r) for r in table_data) if rows>0 else 0
                if rows>0 and cols>0:
                    tbl = doc.add_table(rows, cols)
                    tbl.style = 'Table Grid'
                    for i,r in enumerate(table_data):
                        for j,c in enumerate(r):
                            if j<cols: tbl.cell(i,j).text = c
            in_table=False; table_data=[]
        
        if not stripped: continue
        
        # 2. Gestione Titoli (#, ##)
        if stripped.startswith('#'):
            level = min(stripped.count('#'), 3)
            clean_text = stripped.lstrip('#').strip()
            doc.add_heading(clean_text, level=level)
            
        # 3. Gestione Liste (- o *)
        elif stripped.startswith('- ') or stripped.startswith('* '):
            doc.add_paragraph(stripped[2:], style='List Bullet')
            
        # 4. Paragrafi normali
        else:
            p = doc.add_paragraph(stripped)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def create_zip(docs_dict, sanitizer):
    """Crea lo ZIP finale con i documenti Word"""
    buf = BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as z:
        for name, data in docs_dict.items():
            doc = Document()
            
            # Titolo Documento
            titolo_doc = data.get("titolo", name)
            doc.add_heading(titolo_doc, 0)
            
            # Contenuto (Restore privacy -> Parse Markdown -> Word)
            raw_content = data.get("contenuto", "")
            real_content = sanitizer.restore(raw_content)
            parse_markdown_pro(doc, real_content)
            
            # Salvataggio in memoria
            b = BytesIO()
            doc.save(b)
            z.writestr(f"{name}.docx", b.getvalue())
    
    buf.seek(0)
    return buf
